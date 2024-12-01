from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

with open("guests.txt") as file:
    guests = file.readlines()

guests = [name.strip() for name in guests]

for guest in guests:
    header = doc.add_paragraph("It would be a pleasure to have the company of")
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER

    name = doc.add_paragraph(guest)
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name.runs[0]
    name_run.bold = True
    name_run.font.size = Pt(16)

    details = doc.add_paragraph("at 11010 Memory Lane on the Evening of")
    details.alignment = WD_ALIGN_PARAGRAPH.CENTER

    date = doc.add_paragraph("April 1st")
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date.runs[0]
    date_run.bold = True

    time = doc.add_paragraph("at 7 o'clock")
    time.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if guest != guests[-1]:
        doc.add_page_break()

doc.save("invitations.docx")
